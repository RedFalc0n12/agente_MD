"""
document_generator.py

Responsável por CRIAR os arquivos de saída do agente:
  - gerar um .docx a partir de um texto (título + parágrafos)
  - opcionalmente converter esse .docx para .pdf usando o Gotenberg
    (serviço que você já tem configurado em GOTENBERG_URL, no config.py)

Não depende do Drive nem do LLM — só recebe texto pronto e devolve
caminho(s) de arquivo local(is).
"""

import os
import httpx
from docx import Document

from src.config import settings


def criar_docx(titulo: str, conteudo: str, caminho_saida: str) -> str:
    """
    Cria um .docx simples: título como Heading 1 e o restante do texto
    dividido em parágrafos (quebra por linha em branco).

    `conteudo` é texto puro (o que o LLM gerou). Cada bloco separado por
    uma linha em branco vira um parágrafo novo.
    """
    doc = Document()
    if titulo:
        doc.add_heading(titulo, level=1)

    for bloco in conteudo.split("\n\n"):
        bloco = bloco.strip()
        if bloco:
            doc.add_paragraph(bloco)

    os.makedirs(os.path.dirname(caminho_saida) or ".", exist_ok=True)
    doc.save(caminho_saida)
    return caminho_saida


async def converter_docx_para_pdf(caminho_docx: str, caminho_pdf_saida: str) -> str:
    """
    Converte um .docx em .pdf chamando o Gotenberg (rota
    /forms/libreoffice/convert). Requer o Gotenberg rodando e acessível
    em settings.GOTENBERG_URL (ex: via docker).

    Levanta exceção se o Gotenberg não responder OK — quem chamar deve
    tratar isso (ex: devolver erro 502 pro agente/usuário).
    """
    url = f"{settings.GOTENBERG_URL.rstrip('/')}/forms/libreoffice/convert"
    nome_arquivo = os.path.basename(caminho_docx)

    async with httpx.AsyncClient(timeout=60.0) as client:
        with open(caminho_docx, "rb") as f:
            files = {"files": (nome_arquivo, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            resposta = await client.post(url, files=files)

    resposta.raise_for_status()

    os.makedirs(os.path.dirname(caminho_pdf_saida) or ".", exist_ok=True)
    with open(caminho_pdf_saida, "wb") as f:
        f.write(resposta.content)

    return caminho_pdf_saida


async def gerar_documento(
    titulo: str,
    conteudo: str,
    caminho_base: str,
    formato: str = "pdf",
) -> str:
    """
    Ponto único de entrada: gera o .docx e, se formato == "pdf",
    converte pro pdf e apaga o .docx intermediário.

    `caminho_base` deve vir SEM extensão, ex: "/tmp/abc/relatorio".
    Retorna o caminho final do arquivo gerado.
    """
    formato = formato.lower().strip()
    if formato not in ("docx", "pdf"):
        raise ValueError(f"Formato não suportado: {formato!r}. Use 'docx' ou 'pdf'.")

    caminho_docx = f"{caminho_base}.docx"
    criar_docx(titulo, conteudo, caminho_docx)

    if formato == "docx":
        return caminho_docx

    caminho_pdf = f"{caminho_base}.pdf"
    await converter_docx_para_pdf(caminho_docx, caminho_pdf)
    os.remove(caminho_docx)
    return caminho_pdf